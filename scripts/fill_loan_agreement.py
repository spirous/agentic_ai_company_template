#!/usr/bin/env python3
"""
Loan agreement filler — Agentic Company Workspace

Fills a Word (.docx) loan agreement template by direct run-level manipulation.
The templates contain no {{PLACEHOLDER}} tokens — blanks are tab characters (\t)
and en-space sequences ( ). Each fill point is located by paragraph anchor
text, then the specific run object is replaced.

Template filling is pure Python — no AI, no network, nothing leaves the laptop.

Usage:
  python3 scripts/fill_loan_agreement.py --company "MIT" --type standard --fields fields.json
  python3 scripts/fill_loan_agreement.py --company "MIT" --type ds --fields fields.json

Agreement types:
  standard  Borrower handles all shipping (§7: "None")
  ds        Your company arranges transport both ways (§7: full transport/insurance terms)

Required fields in the JSON file (see REQUIRED_FIELDS below for full list):
  BORROWER_NAME          Full legal name of borrower organisation
  BORROWER_ADDRESS       Full postal address (single line)
  COMPANY_CONTACT        Your contact name + role
  BORROWER_CONTACT       Name + role of borrower contact person
  ARTICLE_DESIGNATION    Product name / model
  SERIAL_NUMBER          Serial number, or "TBD"
  VALUE_CHF              Value in CHF, e.g. "15000"
  REMARKS                Remarks, or "—"
  TOTAL_VALUE_CHF        Total value in CHF
  START_DATE             Loan start date, e.g. "2026-07-15"
  END_DATE               Loan end date, e.g. "2026-10-15"
  AGREEMENT_DATE         Signing date, e.g. "15 July 2026"
  BORROWER_PLACE_DATE    Borrower's place and date, e.g. "Basel, 15 July 2026"
  COMPANY_SIGNER_NAME    Full name of your company's signatory
  COMPANY_SIGNER_TITLE   Title of your company's signatory
  BORROWER_SIGNER_NAME   Full name of borrower signatory
  BORROWER_SIGNER_TITLE  Title of borrower signatory

Optional:
  BORROWER_DELIVERY_ADDRESS  If different from postal address
  COMPANY_SIGNING_CITY       Your company's city for the signing block (default: "[YOUR CITY]")

Output:
  work/loan-agreements/active/YYYYMMDD_Loan_Agreement_Product_YourCompany-Borrower.docx
"""

import os
import sys
import json
import argparse
import re
import zipfile
from datetime import date

PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LEGAL_DIR   = os.path.join(PROJECT_DIR, "engine", "legal", "workflows", "loan-agreements")
TEMPLATES   = {
    "standard": os.path.join(LEGAL_DIR, "templates", "loan_agreement_standard.docx"),
    "ds":       os.path.join(LEGAL_DIR, "templates", "loan_agreement_ds.docx"),
}
ACTIVE_DIR = os.path.join(LEGAL_DIR, "active")

REQUIRED_FIELDS = [
    "BORROWER_NAME",
    "BORROWER_ADDRESS",
    "COMPANY_CONTACT",
    "BORROWER_CONTACT",
    "ARTICLE_DESIGNATION",
    "SERIAL_NUMBER",
    "VALUE_CHF",
    "REMARKS",
    "TOTAL_VALUE_CHF",
    "START_DATE",
    "END_DATE",
    "AGREEMENT_DATE",
    "BORROWER_PLACE_DATE",
    "COMPANY_SIGNER_NAME",
    "COMPANY_SIGNER_TITLE",
    "BORROWER_SIGNER_NAME",
    "BORROWER_SIGNER_TITLE",
]


# ── Paragraph helpers ────────────────────────────────────────────────────────

def _find_para(doc, anchor):
    """Return first paragraph whose stripped text starts with anchor."""
    for p in doc.paragraphs:
        if p.text.strip().startswith(anchor):
            return p
    return None


def _clear_block(runs, start, end, value):
    """Set runs[start].text = value and blank out runs[start+1 .. end-1]."""
    runs[start].text = value
    for i in range(start + 1, end):
        runs[i].text = ''


# ── Fill logic ───────────────────────────────────────────────────────────────

def fill_document(doc, f):
    """Apply all fill points to the Document object in-place.

    NOTE: This function targets the specific paragraph structure of the
    bundled Word templates. If you use a different template, audit the
    paragraph indices and run positions against your own document.
    """

    # Borrower name + address
    # The template has two empty paragraphs immediately before
    # '- ("the Borrower") –'. Fill those empty paragraphs, leave that line untouched.
    borrower_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('- ("the B'):
            borrower_idx = i
            break

    if borrower_idx is not None and borrower_idx >= 2:
        name_para = doc.paragraphs[borrower_idx - 2]
        name_para.add_run(f'{f["BORROWER_NAME"]}, {f["BORROWER_ADDRESS"]}')
        if f.get('BORROWER_DELIVERY_ADDRESS'):
            delivery_para = doc.paragraphs[borrower_idx - 1]
            delivery_para.add_run(f['BORROWER_DELIVERY_ADDRESS'])
    else:
        print('  ⚠  Borrower section not found')

    # Your company contact (single-tab row)
    p = _find_para(doc, '- at [YOUR COMPANY]:')
    if not p:
        p = _find_para(doc, '- at DECTRIS:')  # fallback for adapted templates
    if p:
        p.runs[-1].text = '\t' + f['COMPANY_CONTACT']
    else:
        print('  ⚠  Company contact paragraph not found')

    # Borrower contact (single-tab row)
    p = _find_para(doc, '- at Borrower:')
    if p:
        p.runs[-1].text = '\t' + f['BORROWER_CONTACT']
    else:
        print('  ⚠  Borrower contact paragraph not found')

    # Equipment table
    equipment = {
        'Article designation': ('ARTICLE_DESIGNATION', True),   # single-tab row
        'Serial number':       ('SERIAL_NUMBER',       False),
        'Value (CHF)':         ('VALUE_CHF',           False),
        'Remarks':             ('REMARKS',             False),
        'Total value: CHF':    ('TOTAL_VALUE_CHF',     False),
    }
    for anchor, (key, single_tab) in equipment.items():
        p = _find_para(doc, anchor)
        if p:
            p.runs[-1].text = ('\t' + f[key]) if single_tab else f[key]
        else:
            print(f'  ⚠  Equipment row "{anchor}" not found')

    # Lending period dates
    p = _find_para(doc, 'for the lending period')
    if p:
        runs = p.runs
        runs[7].text = ' ' + f['START_DATE']
        runs[8].text = ' '
        runs[13].text = f['END_DATE'] + '.'
    else:
        print('  ⚠  Lending period paragraph not found')

    # Signing block — Place/Date
    signing_city = f.get('COMPANY_SIGNING_CITY', '[YOUR CITY]')
    p = _find_para(doc, 'Place, Date:')
    if p:
        runs = p.runs
        runs[7].text = signing_city + ', ' + f['AGREEMENT_DATE']
        _clear_block(runs, 15, 20, f['BORROWER_PLACE_DATE'])
    else:
        print('  ⚠  Place/Date signing row not found')

    # Signing block — Names
    p = _find_para(doc, 'Name of signer:')
    if p:
        runs = p.runs
        _clear_block(runs, 6, 11, f['COMPANY_SIGNER_NAME'])
        _clear_block(runs, 18, 23, f['BORROWER_SIGNER_NAME'])
    else:
        print('  ⚠  Name of signer row not found')

    # Signing block — Titles
    p = _find_para(doc, 'Title of signer:')
    if p:
        runs = p.runs
        _clear_block(runs, 6, 11, f['COMPANY_SIGNER_TITLE'])
        _clear_block(runs, 18, 23, f['BORROWER_SIGNER_TITLE'])
    else:
        print('  ⚠  Title of signer row not found')


# ── Post-save XML cleanup ────────────────────────────────────────────────────

def strip_tracked_changes_and_controls(docx_path):
    """
    Remove from the saved .docx:
      - <w:ins> tracked-change insertions (placeholder text left by previous editors,
        visible when Track Changes is on)
      - <w:del> tracked-change deletions
      - <w:sdt> date-picker content controls ("Click or tap to enter a date.")

    Works by patching document.xml inside the zip in-place.
    """
    tmp_path = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin, \
         zipfile.ZipFile(tmp_path, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                xml = data.decode('utf-8')
                xml = re.sub(r'<w:ins\b[^>]*>.*?</w:ins>', '', xml, flags=re.DOTALL)
                xml = re.sub(r'<w:del\b[^>]*>.*?</w:del>', '', xml, flags=re.DOTALL)
                xml = re.sub(r'<w:sdt>.*?</w:sdt>', '', xml, flags=re.DOTALL)
                data = xml.encode('utf-8')
            zout.writestr(item, data)
    os.replace(tmp_path, docx_path)


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Fill a loan agreement template')
    parser.add_argument('--company', required=True,
                        help='Recipient company/institute name (used for output filename)')
    parser.add_argument('--type', required=True, choices=['standard', 'ds'],
                        help='standard = borrower handles shipping; ds = your company handles shipping')
    parser.add_argument('--fields', required=True,
                        help='Path to JSON file with field values')
    args = parser.parse_args()

    try:
        from docx import Document
    except ImportError:
        print('❌  python-docx not installed. Run: pip3 install python-docx')
        sys.exit(1)

    template_path = TEMPLATES[args.type]
    if not os.path.exists(template_path):
        print(f'❌  Template not found: {template_path}')
        print(f'    Add your Word template to: {template_path}')
        sys.exit(1)

    if not os.path.exists(args.fields):
        print(f'❌  Fields file not found: {args.fields}')
        sys.exit(1)

    with open(args.fields, 'r', encoding='utf-8') as fh:
        fields = json.load(fh)

    fields.setdefault('AGREEMENT_DATE', date.today().strftime('%-d %B %Y'))

    missing = [k for k in REQUIRED_FIELDS if k not in fields]
    if missing:
        print(f'⚠️   Missing required fields: {", ".join(missing)}')
        print('    Add them to the JSON file and re-run.')
        sys.exit(1)

    print(f'📄  Filling {args.type} agreement for {args.company}…')
    doc = Document(template_path)
    fill_document(doc, fields)

    os.makedirs(ACTIVE_DIR, exist_ok=True)
    company_slug = re.sub(r'[^a-zA-Z0-9]+', '_', args.company).strip('_')
    product_slug = re.sub(r'[^a-zA-Z0-9]+', '_', fields['ARTICLE_DESIGNATION']).strip('_')
    datestamp = date.today().strftime('%Y%m%d')
    filename = f'{datestamp}_Loan_Agreement_{product_slug}_YourCompany-{company_slug}.docx'
    output_path = os.path.join(ACTIVE_DIR, filename)
    doc.save(output_path)
    strip_tracked_changes_and_controls(output_path)
    print(f'✅  Saved: {output_path}')


if __name__ == '__main__':
    main()
